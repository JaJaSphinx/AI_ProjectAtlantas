import os 
import pandas as pd
from docx import Document
from docx2pdf import convert
from docx.shared import Inches
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
import win32com.client


from test_ollama import ai_ollama


class create_word():
    def create_word_document(headers, df_lease, df_title):
        df_leasefiltered = df_lease.drop(columns=['File Path'])
        df_titlefiltered = df_title.drop(columns=['File Path'])
        doc = Document()
        header_paragraph = doc.add_paragraph("Lease")
        header_paragraph.style = doc.styles['Heading 2']
        # Add a table with header row
        table = doc.add_table(rows=1 + len(df_leasefiltered), cols=len(df_leasefiltered.columns))
        table.style = 'Table Grid'

        # Add headers
        for col_idx, column_name in enumerate(df_leasefiltered.columns):
            table.cell(0, col_idx).text = str(column_name)

        # Add data rows
        for row_idx, row in enumerate(df_leasefiltered.itertuples(index=False), start=1):
            for col_idx, value in enumerate(row):
                table.cell(row_idx, col_idx).text = str(value)


        header_paragraph = doc.add_paragraph("Title")
        header_paragraph.style = doc.styles['Heading 2']
        # Add a table with header row
        table_2 = doc.add_table(rows=1 + len(df_titlefiltered), cols=len(df_titlefiltered.columns))
        table_2.style = 'Table Grid'

        # Add headers
        for col_idx, column_name in enumerate(df_titlefiltered.columns):
            table_2.cell(0, col_idx).text = str(column_name)

        # Add data rows
        for row_idx, row in enumerate(df_titlefiltered.itertuples(index=False), start=1):
            for col_idx, value in enumerate(row):
                table_2.cell(row_idx, col_idx).text = str(value)

        doc.save("output.docx")

class main():
    #Pre-declare params 
    data: list = []
    lease_data: list = []
    title_data: list = []
    lease_id: int = 1
    title_id: int = 1
    id: int = 1
    
    file_path: str = rf"C:\Users\a-lhat\Desktop\Projects\API\project_atlantas\MainFolder"
    headers = ["Folder", "File Name", "Date"]
    #Loop filepath 
    for dirpath, dirnames, filenames in os.walk(file_path):
    
        if "Lease" in dirpath:
            current_section = "Lease"
        elif "Title" in dirpath:
            current_section = "Title"

        for filename in filenames:
            if current_section == "Lease":
                foldername = os.path.basename(dirpath) or os.path.basename(os.path.dirname(dirpath))
                if 'historic' in foldername.lower(): #Ignore any files inside Historic Documents Folder 
                    continue
                full_path = os.path.join(dirpath, filename)
                lease_data.append({"ID": lease_id, 
                            "Folder": foldername,
                            "File Name": filename,
                            "File Path": full_path})
                lease_id += 1
            elif current_section == "Title":
                foldername = os.path.basename(dirpath) or os.path.basename(os.path.dirname(dirpath))
                if 'historic' in foldername.lower(): #Ignore any files inside Historic Documents Folder 
                    continue
                full_path = os.path.join(dirpath, filename)
                title_data.append({"ID": title_id, 
                            "Folder": foldername,
                            "File Name": filename,
                            "File Path": full_path})
                title_id += 1  

    df_lease = pd.DataFrame(lease_data)
    print(df_lease.head())

    df_title = pd.DataFrame(title_data)
    print(df_title.head())    
            
    
    run_ai_ollama_lease = ai_ollama.run_ai(df_lease['File Name'])
    print(run_ai_ollama_lease)
    for lease_date in run_ai_ollama_lease:
        for lease_id in lease_data:
            if lease_date['ID'] == lease_id['ID']:
                lease_id['Date'] = lease_date['Date']
    df_lease = pd.DataFrame(lease_data)
    print(df_lease.head())

    run_ai_ollama_title = ai_ollama.run_ai(df_title['File Name'])
    print(run_ai_ollama_title)
    for title_date in run_ai_ollama_title:
        for title_id in title_data:
            if title_date['ID'] == title_id['ID']:
                title_id['Date'] = title_date['Date']
    df_title = pd.DataFrame(title_data)
    print(df_title.head())

    create = create_word.create_word_document(headers, df_lease, df_title)

    # #Data Frame
    # df = pd.DataFrame(data)
    # print(df.head())
    # run_ai_ollama = ai_ollama.run_ai(df['File Name'])
    # #print(run_ai_ollama)

    # for date in run_ai_ollama:
    #     for i_d in data:
    #         if date['ID'] == i_d['ID']:
    #             i_d['Date'] = date['Date']
    # df = pd.DataFrame(data)
    # print(df.head())
    # #Seperate the dataframe with 'Lease' and 'Title' to enable different tables to be created in Word document 
    # #Probably worth creating 2 different dataframes? Then allow the create_word_document to create two tables using df_1 and df_2 
    # create = create_word.create_word_document(headers, df)

if __name__ == "__main__":
    main()